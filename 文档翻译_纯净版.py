import docx
import requests
import time
import os
import re
import shutil
import tempfile
import zipfile
from lxml import etree
import gspread
from oauth2client.service_account import ServiceAccountCredentials

# API Configuration (using same proxy as 4.0 version)
API_URL = ""
API_KEY = ""
MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"

# Target languages  
LANGUAGES = {
    "German": "DE",
    "French": "FR",
    "Spanish": "ES",
}

class TranslationManager:
    def __init__(self):
        self.translation_memory = {}
        self.terminology_db = {}
        self.consecutive_failures = 0
        self.total_attempts = 0
        self.total_successes = 0
        self.last_api_call_time = 0  # Track last API call for rate limiting
    
    def load_terminology(self, sheet_url, source_lang_col, target_lang_col, language_code):
        """Load terminology from Google Sheets"""
        max_retries = 3
        retry_delay = 2
        
        for attempt in range(max_retries):
            try:
                print(f"  → Loading terminology for {language_code} (attempt {attempt + 1}/{max_retries})")
                
                scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
                credentials = ServiceAccountCredentials.from_json_keyfile_name(
                    r'C:\\Users\\admin\\Desktop\\多语种说明书翻译\\Extract_glossary.json', scope)
                gc = gspread.authorize(credentials)
                
                sheet_id = sheet_url.split('/d/')[1].split('/')[0]
                worksheet = gc.open_by_key(sheet_id).sheet1
                data = worksheet.get_all_records()
                
                if not data or source_lang_col not in data[0].keys() or target_lang_col not in data[0].keys():
                    print(f"  ⚠ No valid terminology data found for {language_code}")
                    return
                
                term_dict = {str(row[source_lang_col]).strip().lower(): str(row[target_lang_col])
                            for row in data if row[source_lang_col] and row[target_lang_col]}
                
                self.terminology_db[language_code] = term_dict
                print(f"  ✓ Loaded {len(term_dict)} terms for {language_code}")
                return
                
            except Exception as e:
                if attempt < max_retries - 1:
                    print(f"  ⚠ Failed to load terminology: {str(e)[:100]}...")
                    print(f"  ⏳ Retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                    retry_delay *= 1.5
                else:
                    print(f"  ✗ Failed to load terminology after {max_retries} attempts: {str(e)[:100]}...")
                    print(f"  ⚠ Continuing without terminology for {language_code}")
                    return
    
    def apply_terminology(self, text, language_code):
        """Apply terminology to text"""
        if language_code not in self.terminology_db or not self.terminology_db[language_code]:
            return text
        
        terms = sorted(self.terminology_db[language_code].keys(), key=len, reverse=True)
        
        def replace_term(match):
            matched_term = match.group(0)
            replacement = self.terminology_db[language_code][match.group(0).lower()]
            
            if matched_term.islower():
                return replacement.lower()
            elif matched_term.isupper():
                return replacement.upper()
            elif matched_term[0].isupper():
                return replacement[0].upper() + replacement[1:]
            return replacement
        
        for term in terms:
            pattern = r'\b' + re.escape(term) + r'\b'
            text = re.sub(pattern, replace_term, text, flags=re.IGNORECASE)
        
        return text
    
    def collect_context(self, text, language_code, max_examples=3, max_length=300):
        """Collect translation context for similar content"""
        if not self.translation_memory:
            return None
        
        words = re.findall(r'\b\w{4,}\b', text.lower())
        if not words:
            return None
        
        matched_entries = []
        for source, translation in self.translation_memory.items():
            score = sum(1 for word in words if word in source.lower())
            if score > 0:
                matched_entries.append((source, translation, score))
        
        matched_entries.sort(key=lambda x: x[2], reverse=True)
        matched_entries = matched_entries[:max_examples]
        
        if not matched_entries:
            return None
        
        context = []
        for source, translation, _ in matched_entries:
            if len(source) > max_length:
                source = source[:max_length] + "..."
            if len(translation) > max_length:
                translation = translation[:max_length] + "..."
            
            context.append(f"EN: {source}\n{language_code}: {translation}")
        
        return "\n\n".join(context)
    
    def verify_translation(self, translated_text, original_text):
        """Clean and verify translation"""
        english_indicators = [
            "I'm sorry", "I apologize", "Here is the translation",
            "Translated text", "Please note", "I cannot", "I would",
            "Dear Valued Customer", "Best regards", "The Customer Service Team"
        ]
        
        for phrase in english_indicators:
            if phrase.lower() in translated_text.lower():
                clean_lines = [line for line in translated_text.split('\n') 
                              if not any(indicator.lower() in line.lower() 
                                        for indicator in english_indicators)]
                translated_text = '\n'.join(clean_lines)
        
        if not translated_text.strip() and original_text.strip():
            return original_text
        
        if not re.search(r'[a-zA-Z]', original_text) and len(original_text.strip()) < 3:
            return original_text
            
        return translated_text
    
    def translate_text(self, text, target_language, language_code, context=None, is_footnote=False):
        """Translate text using API with context and terminology support"""
        if not text.strip():
            return ""
        
        # Skip Chinese content
        chinese_char_count = len(re.findall(r'[\u4e00-\u9fff]', text))
        text_length = len(text.strip())
        if chinese_char_count > 0 and chinese_char_count / text_length > 0.1:
            return text

        # Check translation memory
        memory_key = text.strip().lower()
        if memory_key in self.translation_memory:
            return self.translation_memory[memory_key]
        
        # Show translation progress (simplified for clean version)
        text_preview = text[:50] + "..." if len(text) > 50 else text
        print(f"Translating: {text_preview}")
        
        # Enhanced consecutive failures check with rate limiting protection
        if self.consecutive_failures >= 5:  # Reduced threshold from 10 to 5 for faster intervention
            print(f"\n⚠️ Warning: {self.consecutive_failures} consecutive translation failures detected.")
            print(f"Network statistics: {self.total_successes}/{self.total_attempts} successful")
            print("This may indicate API rate limiting or connectivity issues.")
            
            # Add automatic cooling-off period for frequent failures
            if self.consecutive_failures >= 8:
                cooling_period = min(30, self.consecutive_failures * 2)  # Cap at 30 seconds
                print(f"Implementing {cooling_period}s cooling-off period to avoid rate limits...")
                time.sleep(cooling_period)
            
            user_choice = input("Continue translation? (y/n): ").strip().lower()
            if user_choice != 'y':
                print("Translation stopped by user.")
                return text
            else:
                self.consecutive_failures = 0  # Reset counter if user chooses to continue
        
        # Extract terminology for context
        potential_terms = []
        if language_code in self.terminology_db:
            for term in self.terminology_db[language_code].keys():
                if re.search(r'\b' + re.escape(term) + r'\b', text, re.IGNORECASE):
                    target_term = self.terminology_db[language_code][term]
                    potential_terms.append(f"{term} -> {target_term}")
        
        try:
            if len(text.strip()) < 3 and not re.search(r'[a-zA-Z]', text):
                return text
            
            # Build prompt
            if is_footnote:
                sys_prompt = 'You are a translation engine specialized in footnotes. Translate the footnote text to the target language maintaining all formatting and academic/reference style. Return ONLY the translated text with no explanations, no English, and no comments.'
                user_prompt = f'Translate the following footnote text to {target_language}. Maintain the scholarly and reference tone typical of footnotes. Return ONLY the translated footnote content. Keep all symbols, punctuation, and formatting exactly as they appear.'
            else:
                sys_prompt = f'You are a professional translation engine. Translate text from English to {target_language} maintaining all formatting. Return ONLY the translated text with no explanations, no English text, and no comments. Never apologize or explain your translation.'
                if target_language == "Spanish":
                    user_prompt = f'Translate the following English text to {target_language}. Use neutral Spanish that is appropriate for technical/marketing documentation. Return ONLY the Spanish translation. Keep all symbols, punctuation, and formatting exactly as they appear. Do not add any explanations, English text, or comments before or after the translation.'
                else:
                    user_prompt = f'Translate the following text to {target_language}. Return ONLY the translated content. Keep all symbols, punctuation, and formatting exactly as they appear. Do not add any explanations, English text, or comments before or after the translation.'
            
            user_prompt += "\n\nIMPORTANT: If the text contains only symbols, formatting characters, or no text at all (like '----', '***', etc.), do not translate or explain anything - just return those exact symbols."
            
            if potential_terms:
                user_prompt += f"\n\nIMPORTANT: Use the following terminology consistently:\n" + "\n".join(potential_terms)
            
            if context:
                user_prompt += f"\n\nFor consistency, here are some previous translations:\n{context}\n\n"
            
            user_prompt += f"\nText to translate:\n{text}"
            
            # Request data
            data = {
                "model": MODEL_ID,
                "messages": [
                    {
                        "role": "user",
                        "content": [{"type": "text", "text": user_prompt}],
                    }
                ],
                "system": sys_prompt
            }
            
            headers = {
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {API_KEY}'
            }
            
            # Rate limiting: Ensure minimum interval between API calls
            min_interval = 1.0  # Minimum 1 second between API calls
            time_since_last_call = time.time() - self.last_api_call_time
            if time_since_last_call < min_interval:
                wait_time = min_interval - time_since_last_call
                print(f"  ⏳ Rate limiting: waiting {wait_time:.1f}s...")
                time.sleep(wait_time)
            
            # API request with enhanced retry logic (adopted from 4.0 version)
            max_retries = 3  # Reduced from 5 to minimize API call frequency
            retry_delay = 5  # Increased from 3 to reduce request frequency
            retry_count = 0
            self.total_attempts += 1

            while retry_count < max_retries:
                try:
                    print(f"  → API request (attempt {retry_count + 1}/{max_retries})")
                    self.last_api_call_time = time.time()  # Record API call time
                    response = requests.post(API_URL, headers=headers, json=data, timeout=60)
                    
                    if response.status_code == 200:
                        result = response.json()
                        translated_text = result["choices"][0]["message"]["content"]
                        print(f"  ✓ Translation successful")
                        
                        # Clean and verify translation
                        translated_text = self.verify_translation(translated_text, text)
                        explanation_patterns = [
                            r'^(I\'m sorry|I apologize|Sorry|Note|Please note).*?\n\n',
                            r'\n\n(I\'m sorry|I apologize|Sorry|Note|Please note).*?$',
                            r'^(Here is|Here\'s|The following is|This is) the translation.*?\n\n',
                            r'^Translated text:.*?\n\n'
                        ]
                        
                        for pattern in explanation_patterns:
                            translated_text = re.sub(pattern, '', translated_text, flags=re.IGNORECASE | re.DOTALL)

                        # Apply terminology
                        translated_text = self.apply_terminology(translated_text, language_code)
                        
                        # Store in memory
                        self.translation_memory[memory_key] = translated_text
                        
                        # Update success counters
                        self.total_successes += 1
                        self.consecutive_failures = 0  # Reset consecutive failures on success
                        
                        return translated_text
                    
                    elif response.status_code == 502:
                        # Handle 502 errors with longer delays (from 4.0 version)
                        retry_count += 1
                        print(f"  ⚠ Server error (502). Retry attempt {retry_count}/{max_retries} in {retry_delay:.1f}s...")
                        if retry_count < max_retries:
                            time.sleep(retry_delay)
                            retry_delay *= 1.5  # More aggressive delay increase for 502 errors
                        continue
                    
                    elif response.status_code == 429:
                        retry_count += 1
                        print(f"  ⚠ Rate limit exceeded. Retrying in {retry_delay * 2:.1f}s...")
                        if retry_count < max_retries:
                            time.sleep(retry_delay * 2)  # Longer delay for rate limits
                            retry_delay *= 1.5
                        continue
                        
                    elif response.status_code == 401:
                        print(f"  ✗ Authentication failed (401). Check API key.")
                        self.consecutive_failures += 1
                        return text
                        
                    else:
                        # Handle other HTTP errors with shorter delays (from 4.0 version)
                        retry_count += 1
                        shorter_delay = retry_delay / 2
                        print(f"  ⚠ HTTP {response.status_code} error. Retry attempt {retry_count}/{max_retries} in {shorter_delay:.1f}s...")
                        if retry_count < max_retries:
                            time.sleep(shorter_delay)
                        continue
                        
                except requests.exceptions.ConnectionError as e:
                    retry_count += 1
                    print(f"  ✗ Network connection failed: {str(e)[:100]}...")
                    if retry_count < max_retries:
                        print(f"  ⏳ Retrying in {retry_delay:.1f}s...")
                        time.sleep(retry_delay)
                        retry_delay *= 1.5
                    continue
                    
                except requests.exceptions.Timeout as e:
                    retry_count += 1
                    print(f"  ✗ Request timeout: {str(e)[:100]}...")
                    if retry_count < max_retries:
                        print(f"  ⏳ Retrying in {retry_delay:.1f}s...")
                        time.sleep(retry_delay)
                        retry_delay *= 1.3
                    continue
                    
                except requests.exceptions.RequestException as e:
                    retry_count += 1
                    print(f"  ✗ Request error: {str(e)[:100]}...")
                    if retry_count < max_retries:
                        print(f"  ⏳ Retrying in {retry_delay:.1f}s...")
                        time.sleep(retry_delay)
                        retry_delay *= 1.5  # Increased from 1.4 to reduce request frequency
                    continue
                    
                except Exception as e:
                    retry_count += 1
                    print(f"  ✗ Unexpected error: {str(e)[:100]}...")
                    if retry_count < max_retries:
                        print(f"  ⏳ Retrying in {retry_delay:.1f}s...")
                        time.sleep(retry_delay)
                        retry_delay *= 1.5  # Increased from 1.2
                    continue

            # All retries failed - enhanced error reporting
            print(f"  ❌ Translation failed after {max_retries} attempts for text: '{text[:50]}...'. Using original text.")
            self.consecutive_failures += 1
            return text

        except Exception:
            return text
    
    def clear_memory(self):
        """Clear translation memory"""
        self.translation_memory = {}
        self.consecutive_failures = 0
        self.total_attempts = 0
        self.total_successes = 0
        self.last_api_call_time = 0  # Reset API call timing

class DocumentProcessor:
    def __init__(self, translator):
        self.translator = translator
    
    def capture_run_properties(self, run):
        """Capture all run properties with robust color handling"""
        properties = {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "strike": run.strike if hasattr(run, "strike") else None,
            "highlight_color": None,
            "color": None,
            "name": run.font.name if hasattr(run.font, "name") else None,
            "size": run.font.size if hasattr(run.font, "size") else None,
            "subscript": run.font.subscript if hasattr(run.font, "subscript") else None,
            "superscript": run.font.superscript if hasattr(run.font, "superscript") else None,
            "all_caps": run.font.all_caps if hasattr(run.font, "all_caps") else None,
            "double_strike": run.font.double_strike if hasattr(run.font, "double_strike") else None,
            "emboss": run.font.emboss if hasattr(run.font, "emboss") else None,
            "imprint": run.font.imprint if hasattr(run.font, "imprint") else None,
            "outline": run.font.outline if hasattr(run.font, "outline") else None,
            "rtl": run.font.rtl if hasattr(run.font, "rtl") else None,
            "shadow": run.font.shadow if hasattr(run.font, "shadow") else None,
            "small_caps": run.font.small_caps if hasattr(run.font, "small_caps") else None,
            "snap_to_grid": run.font.snap_to_grid if hasattr(run.font, "snap_to_grid") else None,
            "spec_vanish": run.font.spec_vanish if hasattr(run.font, "spec_vanish") else None
        }

        # Handle highlight color
        try:
            if hasattr(run.font, "highlight_color"):
                highlight = run.font.highlight_color
                if highlight is not None and highlight != 'none' and highlight != 'auto':
                    properties["highlight_color"] = highlight
        except Exception:
            pass

        # Handle font color
        try:
            if hasattr(run.font, "color") and run.font.color is not None:
                if hasattr(run.font.color, "rgb") and run.font.color.rgb is not None:
                    if run.font.color.rgb != 'none' and run.font.color.rgb != 'auto':
                        properties["color"] = run.font.color.rgb
        except Exception:
            pass
            
        return properties

    def apply_run_properties(self, run, properties):
        """Apply run properties with enhanced color handling"""
        run.bold = properties["bold"]
        run.italic = properties["italic"]
        run.underline = properties["underline"]
        
        if properties["strike"] is not None and hasattr(run, "strike"):
            run.strike = properties["strike"]
        
        if properties["name"] is not None:
            run.font.name = properties["name"]
        if properties["size"] is not None:
            run.font.size = properties["size"]
        
        if properties["color"] is not None and hasattr(run.font, "color"):
            try:
                color_value = properties["color"]
                if isinstance(color_value, bytes) or (
                    isinstance(color_value, str) and 
                    color_value not in ('none', 'auto', '')
                ):
                    run.font.color.rgb = color_value
            except Exception:
                try:
                    if hasattr(run.font.color, "_element"):
                        run.font.color._element.clear()
                except:
                    pass
        
        if properties["highlight_color"] is not None and hasattr(run.font, "highlight_color"):
            try:
                hl_value = properties["highlight_color"]
                if hl_value not in ('none', 'auto', ''):
                    run.font.highlight_color = hl_value
            except Exception:
                pass
        
        for attr in ["subscript", "superscript", "all_caps", "double_strike", 
                    "emboss", "imprint", "outline", "rtl", "shadow", 
                    "small_caps", "snap_to_grid", "spec_vanish"]:
            if properties[attr] is not None and hasattr(run.font, attr):
                try:
                    setattr(run.font, attr, properties[attr])
                except Exception:
                    pass
        
        return run

    def process_paragraph(self, paragraph, target_language, language_code):
        """Process and translate paragraph with format preservation"""
        if not paragraph.text.strip():
            return paragraph
        
        # Store original formatting
        runs_formatting = []
        runs_text = []
        
        for run in paragraph.runs:
            runs_text.append(run.text)
            runs_formatting.append(self.capture_run_properties(run))
        
        if not runs_text:
            return paragraph
        
        # Get full paragraph text for translation
        text = paragraph.text
        
        if text.strip():
            context = self.translator.collect_context(text, language_code)
            translated_text = self.translator.translate_text(text, target_language, language_code, context)
            
            if not translated_text or (translated_text == text and len(text.strip()) < 3 and not re.search(r'[a-zA-Z]', text)):
                return paragraph
            
            if translated_text != text:
                paragraph.clear()
                new_run = paragraph.add_run(translated_text)
                
                if runs_formatting:
                    try:
                        self.apply_run_properties(new_run, runs_formatting[0])
                    except Exception:
                        pass
        
        return paragraph
    
    def process_table(self, table, target_language, language_code):
        """Process and translate table content"""
        try:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        try:
                            self.process_paragraph(para, target_language, language_code)
                        except Exception:
                            pass
        except Exception:
            pass
    
    def has_tables(self, doc):
        """Check if document contains tables"""
        return len(doc.tables) > 0
    
    def has_footnotes(self, doc_path):
        """Check if document contains footnotes"""
        try:
            with zipfile.ZipFile(doc_path, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                
                has_footnote_files = any(f in file_list for f in ['word/footnotes.xml', 'word/endnotes.xml'])
                
                if has_footnote_files and 'word/document.xml' in file_list:
                    doc_content = zip_ref.read('word/document.xml').decode('utf-8')
                    has_footnote_refs = ('<w:footnoteReference' in doc_content or 
                                       '<w:endnoteReference' in doc_content)
                    return has_footnote_refs
                
                return False
        except Exception:
            return False

    def process_footnotes_with_merge(self, original_file, translated_file, output_file, target_language, language_code):
        """Extract footnotes from original document, translate and merge into translated document"""
        import re
        
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                original_dir = os.path.join(temp_dir, "original")
                translated_dir = os.path.join(temp_dir, "translated")
                
                os.makedirs(original_dir, exist_ok=True)
                os.makedirs(translated_dir, exist_ok=True)
                
                # Extract documents
                with zipfile.ZipFile(original_file, 'r') as zip_ref:
                    zip_ref.extractall(original_dir)
                
                with zipfile.ZipFile(translated_file, 'r') as zip_ref:
                    zip_ref.extractall(translated_dir)
                
                # Process footnote files
                footnote_files = [
                    ('word/footnotes.xml', 'footnote'),
                    ('word/endnotes.xml', 'endnote')
                ]
                
                total_translations = 0
                
                for file_path, note_type in footnote_files:
                    original_footnote_path = os.path.join(original_dir, file_path)
                    translated_footnote_path = os.path.join(translated_dir, file_path)
                    
                    if not os.path.exists(original_footnote_path):
                        continue
                    
                    try:
                        shutil.copy2(original_footnote_path, translated_footnote_path)
                        
                        parser = etree.XMLParser(strip_cdata=False, recover=True)
                        tree = etree.parse(translated_footnote_path, parser)
                        root = tree.getroot()
                        
                        namespaces = {
                            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        }
                        
                        text_elements = root.xpath('//w:t', namespaces=namespaces)
                        translations_made = 0
                        
                        import re
                        for i, text_elem in enumerate(text_elements):
                            if text_elem.text is None:
                                continue
                                
                            original_text = text_elem.text
                            
                            if (not original_text.strip() or 
                                len(original_text.strip()) < 3 or
                                re.match(r'^[\d\s\.\-_]+$', original_text.strip()) or
                                not re.search(r'[a-zA-Z]', original_text)):
                                continue
                            
                            try:
                                context = self.translator.collect_context(original_text, language_code)
                                translated_text = self.translator.translate_text(
                                    original_text, target_language, language_code, context, is_footnote=True
                                )
                                
                                if translated_text != original_text and translated_text.strip():
                                    text_elem.text = translated_text
                                    translations_made += 1
                                    
                            except Exception:
                                pass
                        
                        if translations_made > 0:
                            print(f"Translated {translations_made} footnote texts.")
                        
                        if translations_made > 0:
                            tree.write(translated_footnote_path, encoding='utf-8', xml_declaration=True, 
                                     pretty_print=False, method='xml')
                            total_translations += translations_made
                        
                    except Exception:
                        continue
                
                # Copy footnote relationship files
                footnote_rels_files = [
                    'word/_rels/footnotes.xml.rels',
                    'word/_rels/endnotes.xml.rels'
                ]
                
                for rel_file in footnote_rels_files:
                    original_rel_path = os.path.join(original_dir, rel_file)
                    translated_rel_path = os.path.join(translated_dir, rel_file)
                    
                    if os.path.exists(original_rel_path):
                        os.makedirs(os.path.dirname(translated_rel_path), exist_ok=True)
                        shutil.copy2(original_rel_path, translated_rel_path)
                
                # Fix footnote references in main document
                original_doc_xml = os.path.join(original_dir, 'word/document.xml')
                translated_doc_xml = os.path.join(translated_dir, 'word/document.xml')
                
                if os.path.exists(original_doc_xml) and os.path.exists(translated_doc_xml):
                    try:
                        with open(original_doc_xml, 'r', encoding='utf-8') as f:
                            original_doc_content = f.read()
                        
                        with open(translated_doc_xml, 'r', encoding='utf-8') as f:
                            translated_doc_content = f.read()
                        
                        footnote_refs = re.findall(r'<w:footnoteReference[^>]*>', original_doc_content)
                        current_refs = re.findall(r'<w:footnoteReference[^>]*>', translated_doc_content)
                        
                        if len(footnote_refs) > len(current_refs):
                            original_paragraphs = re.findall(r'<w:p\b[^>]*>.*?</w:p>', original_doc_content, re.DOTALL)
                            translated_paragraphs = re.findall(r'<w:p\b[^>]*>.*?</w:p>', translated_doc_content, re.DOTALL)
                            
                            updated_content = translated_doc_content
                            refs_inserted = 0
                            
                            for i, orig_para in enumerate(original_paragraphs):
                                if '<w:footnoteReference' in orig_para:
                                    para_refs = re.findall(r'<w:footnoteReference[^>]*>', orig_para)
                                    
                                    if i < len(translated_paragraphs):
                                        trans_para = translated_paragraphs[i]
                                        para_start = updated_content.find(trans_para)
                                        
                                        if para_start != -1:
                                            para_end_tag = '</w:p>'
                                            para_end = updated_content.find(para_end_tag, para_start)
                                            
                                            if para_end != -1:
                                                for ref in para_refs:
                                                    footnote_run = f'<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr>{ref}</w:r>'
                                                    updated_content = (
                                                        updated_content[:para_end] + 
                                                        footnote_run + 
                                                        updated_content[para_end:]
                                                    )
                                                    refs_inserted += 1
                                                    para_end += len(footnote_run)
                            
                            if refs_inserted > 0:
                                with open(translated_doc_xml, 'w', encoding='utf-8') as f:
                                    f.write(updated_content)
                            
                    except Exception:
                        pass
                
                # Repackage document
                if total_translations > 0:
                    with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED, compresslevel=6) as docx_file:
                        for root_dir, _, files in os.walk(translated_dir):
                            for file in files:
                                file_path = os.path.join(root_dir, file)
                                arcname = os.path.relpath(file_path, translated_dir)
                                docx_file.write(file_path, arcname)
                    
                    return True
                else:
                    shutil.copy2(translated_file, output_file)
                    return False
                    
        except Exception:
            shutil.copy2(translated_file, output_file)
            return False

    def has_text_boxes(self, doc_path):
        """Check if document contains text boxes"""
        try:
            with zipfile.ZipFile(doc_path, 'r') as zip_ref:
                xml_files = ['word/document.xml', 'word/header1.xml', 'word/header2.xml', 'word/header3.xml', 
                            'word/footer1.xml', 'word/footer2.xml', 'word/footer3.xml']
                
                for xml_file in xml_files:
                    if xml_file in zip_ref.namelist():
                        content = zip_ref.read(xml_file).decode('utf-8')
                        if '<w:txbxContent>' in content:
                            return True
                return False
        except Exception:
            return False

    def preserve_images(self, input_file, translated_file):
        """Preserve images with exact positioning"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                original_dir = os.path.join(temp_dir, "original")
                translated_dir = os.path.join(temp_dir, "translated")
                
                os.makedirs(original_dir, exist_ok=True)
                os.makedirs(translated_dir, exist_ok=True)
                
                # Extract both documents
                with zipfile.ZipFile(input_file, 'r') as zip_ref:
                    zip_ref.extractall(original_dir)
                
                with zipfile.ZipFile(translated_file, 'r') as zip_ref:
                    zip_ref.extractall(translated_dir)
                
                # Copy media folder
                original_media_dir = os.path.join(original_dir, "word", "media")
                translated_media_dir = os.path.join(translated_dir, "word", "media")
                
                if os.path.exists(original_media_dir):
                    if os.path.exists(translated_media_dir):
                        shutil.rmtree(translated_media_dir)
                    
                    shutil.copytree(original_media_dir, translated_media_dir)
                
                # Process relationships and content types
                self._process_document_relationships(original_dir, translated_dir)
                self._process_document_content_types(original_dir, translated_dir)
                self._process_document_structure(original_dir, translated_dir)
                
                # Repackage document
                output_temp = translated_file + ".temp.docx"
                
                with zipfile.ZipFile(output_temp, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                    for root, _, files in os.walk(translated_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, translated_dir)
                            zip_out.write(file_path, arcname)
                
                # Replace original output
                if os.path.exists(translated_file):
                    os.remove(translated_file)
                os.rename(output_temp, translated_file)
                
        except Exception:
            pass
    
    def _process_document_relationships(self, original_dir, translated_dir):
        """Process document relationships to maintain image references"""
        original_rels_file = os.path.join(original_dir, "word", "_rels", "document.xml.rels")
        translated_rels_file = os.path.join(translated_dir, "word", "_rels", "document.xml.rels")
        
        if not os.path.exists(original_rels_file):
            return
        
        os.makedirs(os.path.dirname(translated_rels_file), exist_ok=True)
        
        try:
            image_relations = {}
            tree = etree.parse(original_rels_file)
            root = tree.getroot()
            
            for rel in root:
                if all(attr in rel.attrib for attr in ['Id', 'Type', 'Target']):
                    rel_id = rel.attrib['Id']
                    rel_type = rel.attrib['Type']
                    target = rel.attrib['Target']
                    
                    if 'image' in rel_type.lower() or target.startswith('media/'):
                        image_relations[rel_id] = {
                            'type': rel_type,
                            'target': target
                        }
            
            if os.path.exists(translated_rels_file):
                tree = etree.parse(translated_rels_file)
                root = tree.getroot()
                
                existing_ids = {rel.attrib['Id'] for rel in root if 'Id' in rel.attrib}
                
                for rel_id, rel_info in image_relations.items():
                    if rel_id not in existing_ids:
                        ns = '{http://schemas.openxmlformats.org/package/2006/relationships}'
                        new_rel = etree.SubElement(root, ns + 'Relationship')
                        new_rel.set('Id', rel_id)
                        new_rel.set('Type', rel_info['type'])
                        new_rel.set('Target', rel_info['target'])
                        new_rel.set('TargetMode', 'Internal')
                
                tree.write(translated_rels_file, encoding='UTF-8', xml_declaration=True)
            else:
                shutil.copy2(original_rels_file, translated_rels_file)
                
        except Exception:
            if os.path.exists(original_rels_file):
                shutil.copy2(original_rels_file, translated_rels_file)
    
    def _process_document_content_types(self, original_dir, translated_dir):
        """Process content types to include image formats"""
        content_types_file = os.path.join(original_dir, "[Content_Types].xml")
        translated_types_file = os.path.join(translated_dir, "[Content_Types].xml")
        
        if not os.path.exists(content_types_file) or not os.path.exists(translated_types_file):
            return
            
        try:
            original_tree = etree.parse(content_types_file)
            original_root = original_tree.getroot()
            
            translated_tree = etree.parse(translated_types_file)
            translated_root = translated_tree.getroot()
            
            existing_extensions = {default.attrib['Extension'] 
                                for default in translated_root.findall(".//{*}Default") 
                                if 'Extension' in default.attrib}
            
            image_extensions = ['png', 'jpeg', 'jpg', 'gif', 'bmp', 'tiff', 'emf', 'wmf']
            modified = False
            
            for default in original_root.findall(".//{*}Default"):
                if 'Extension' in default.attrib:
                    ext = default.attrib['Extension']
                    if ext in image_extensions and ext not in existing_extensions:
                        translated_root.append(default)
                        modified = True
            
            if modified:
                translated_tree.write(translated_types_file, encoding='UTF-8', xml_declaration=True)
                
        except Exception:
            shutil.copy2(content_types_file, translated_types_file)
    
    def _process_document_structure(self, original_dir, translated_dir):
        """Process document structure to position images correctly"""
        original_doc_xml = os.path.join(original_dir, "word", "document.xml")
        translated_doc_xml = os.path.join(translated_dir, "word", "document.xml")
        
        if not os.path.exists(original_doc_xml) or not os.path.exists(translated_doc_xml):
            return
            
        try:
            with open(original_doc_xml, 'r', encoding='utf-8') as f:
                original_content = f.read()
            
            with open(translated_doc_xml, 'r', encoding='utf-8') as f:
                translated_content = f.read()
            
            original_paragraphs = re.findall(r'<w:p\b[^>]*>.*?</w:p>', original_content, re.DOTALL)
            translated_paragraphs = re.findall(r'<w:p\b[^>]*>.*?</w:p>', translated_content, re.DOTALL)
            
            image_paragraphs = []
            for i, para in enumerate(original_paragraphs):
                if '<w:drawing>' in para:
                    img_refs = re.findall(r'r:embed="(rId\d+)"', para)
                    if img_refs:
                        image_paragraphs.append({
                            'index': i,
                            'paragraph': para,
                            'image_refs': img_refs
                        })
            
            missing_images = []
            for img_para in image_paragraphs:
                for ref_id in img_para['image_refs']:
                    if f'r:embed="{ref_id}"' not in translated_content:
                        drawing_match = re.search(
                            f'<w:drawing>.*?r:embed="{ref_id}".*?</w:drawing>', 
                            img_para['paragraph'], 
                            re.DOTALL
                        )
                        if drawing_match:
                            missing_images.append({
                                'ref_id': ref_id,
                                'drawing': drawing_match.group(0),
                                'para_index': img_para['index']
                            })
            
            if missing_images:
                modified_content = translated_content
                original_para_count = len(original_paragraphs)
                translated_para_count = len(translated_paragraphs)
                
                for img_info in missing_images:
                    relative_pos = img_info['para_index'] / original_para_count
                    target_para_index = min(max(0, int(relative_pos * translated_para_count)), 
                                        translated_para_count - 1)
                    
                    if target_para_index < len(translated_paragraphs):
                        target_para = translated_paragraphs[target_para_index]
                        para_pos = modified_content.find(target_para)
                        
                        if para_pos != -1:
                            para_end = para_pos + len(target_para)
                            img_para = f'<w:p><w:r>{img_info["drawing"]}</w:r></w:p>'
                            modified_content = (
                                modified_content[:para_end] + 
                                img_para + 
                                modified_content[para_end:]
                            )
                        else:
                            body_end = modified_content.rfind('</w:body>')
                            if body_end != -1:
                                img_para = f'<w:p><w:r>{img_info["drawing"]}</w:r></w:p>'
                                modified_content = (
                                    modified_content[:body_end] + 
                                    img_para + 
                                    modified_content[body_end:]
                                )
                
                with open(translated_doc_xml, 'w', encoding='utf-8') as f:
                    f.write(modified_content)
                    
        except Exception:
            pass

class DocumentTranslator:
    def __init__(self):
        self.translator = TranslationManager()
        self.processor = DocumentProcessor(self.translator)
    
    def translate_document(self, input_file, output_dir, google_sheet_url=None):
        """Translate document to all target languages"""
        print("Starting document translation...")
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Load terminology if provided
        if google_sheet_url:
            print("Loading terminology...")
            for language_name, language_code in LANGUAGES.items():
                self.translator.load_terminology(
                    google_sheet_url, 
                    "English",
                    language_name,
                    language_code
                )
            print("Terminology loaded.")
        
        # Process each target language
        for language_name, language_code in LANGUAGES.items():
            try:
                print(f"\n=== Translating to {language_name} ===")
                self.translator.clear_memory()
                
                # Prepare output file
                base_name = os.path.splitext(os.path.basename(input_file))[0]
                output_file = os.path.join(output_dir, f"{base_name}_{language_code}.docx")
                
                # Step 1: Load and translate main content
                print("Translating main content...")
                doc = docx.Document(input_file)
                
                # Translate main paragraphs
                paragraph_count = 0
                for para in doc.paragraphs:
                    if para.text.strip():
                        self.processor.process_paragraph(para, language_name, language_code)
                        paragraph_count += 1
                print(f"Main content completed: {paragraph_count} paragraphs translated.")
                
                # Translate tables
                if self.processor.has_tables(doc):
                    print("Translating tables...")
                    table_count = 0
                    for table in doc.tables:
                        self.processor.process_table(table, language_name, language_code)
                        table_count += 1
                    print(f"Tables completed: {table_count} tables translated.")
                
                # Translate headers and footers
                print("Translating headers and footers...")
                for section in doc.sections:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            self.processor.process_paragraph(para, language_name, language_code)
                    
                    for table in section.header.tables:
                        self.processor.process_table(table, language_name, language_code)
                    
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            self.processor.process_paragraph(para, language_name, language_code)
                    
                    for table in section.footer.tables:
                        self.processor.process_table(table, language_name, language_code)
                print("Headers and footers completed.")
                
                # Save intermediate document
                intermediate_file = output_file.replace('.docx', '_temp.docx')
                doc.save(intermediate_file)
                
                current_file = intermediate_file
                
                # Process footnotes
                if self.processor.has_footnotes(input_file):
                    print("Translating footnotes...")
                    footnote_file = current_file.replace('.docx', '_footnotes.docx')
                    footnote_success = self.processor.process_footnotes_with_merge(
                        input_file, current_file, footnote_file, language_name, language_code
                    )
                    if footnote_success:
                        current_file = footnote_file
                        print("Footnotes completed.")
                    else:
                        print("No footnotes to translate.")
                
                # Copy current file to final output
                if current_file != output_file:
                    shutil.copy2(current_file, output_file)
                
                # Preserve images
                print("Processing images...")
                self.processor.preserve_images(input_file, output_file)
                print("Images processed.")
                
                # Clean up intermediate files
                cleanup_files = [intermediate_file]
                if current_file != intermediate_file and current_file != output_file:
                    cleanup_files.append(current_file)
                
                for cleanup_file in cleanup_files:
                    if os.path.exists(cleanup_file):
                        try:
                            os.remove(cleanup_file)
                        except Exception:
                            pass
                
                # Show translation statistics
                success_rate = (self.translator.total_successes / self.translator.total_attempts * 100) if self.translator.total_attempts > 0 else 0
                print(f"Network stats: {self.translator.total_successes}/{self.translator.total_attempts} successful ({success_rate:.1f}%)")
                print(f"✓ {language_name} translation completed: {output_file}")
                
            except Exception as e:
                print(f"✗ Error translating to {language_name}: {e}")
        
        print("\n=== All translations completed ===")
        print(f"Output directory: {output_dir}")

def main():
    """Main function"""
    input_file = r"C:\\Users\\admin\\Desktop\\Selling points\\EN\\Selling Points Text Version-Aqara Camera G100 Select.docx"
    output_dir = r"C:\\Users\\admin\\Desktop\\Selling points\\G100"
    google_sheet_url = "https://docs.google.com/spreadsheets/d/11B4LNWf27Mt_PvqsyKZYmtxeaLmCBPFSQHiiUyY2IC4/edit?gid=0"
    
    print("Document Translation Tool")
    print("=" * 50)
    print(f"Input file: {input_file}")
    print(f"Output directory: {output_dir}")
    print("=" * 50)
    
    if not os.path.exists(input_file):
        print(f"✗ Error: Input file not found: {input_file}")
        return
    
    try:
        translator = DocumentTranslator()
        translator.translate_document(input_file, output_dir, google_sheet_url)
        print("\n🎉 All translations completed successfully!")
    except Exception as e:
        print(f"\n✗ Translation failed: {e}")

if __name__ == "__main__":

    main() 
