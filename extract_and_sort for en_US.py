import openpyxl
import re
from docx import Document
from docx.shared import Pt
import os

def extract_questions_and_answers(excel_file):
    # 打开 Excel 文件
    wb = openpyxl.load_workbook(excel_file)
    sheet = wb.active

    # 在第一行查找 "en-US" 标题所在的列（忽略大小写和前后空格）
    header_row = list(sheet.iter_rows(min_row=1, max_row=1, values_only=True))[0]
    en_us_col = None
    for idx, header in enumerate(header_row):
        if isinstance(header, str) and header.strip().lower() == "en-us":
            en_us_col = idx + 1  # openpyxl 的行列号是 1 开始
            break
    if en_us_col is None:
        raise ValueError("未找到 en-US 列")

    qa_pairs = {}  # 用字典保存：键为题目编号，值为 { 'question': ..., 'answer': ... }

    # 从第二行开始遍历所有行，提取第一列（代码）和 en-US 列（文本）
    for row in sheet.iter_rows(min_row=2, values_only=True):
        code = row[0]  # 第一列，代码
        text = row[en_us_col - 1]  # 因为 row 是 0 索引
        if isinstance(code, str) and isinstance(text, str):
            code = code.strip()
            text = text.strip()
            # 查找包含 "title-数字" 和 "content-数字" 的信息
            title_match = re.search(r'title-(\d+)', code, re.IGNORECASE)
            content_match = re.search(r'content-(\d+)', code, re.IGNORECASE)
            if title_match:
                qid = title_match.group(1)
                if qid not in qa_pairs:
                    qa_pairs[qid] = {'question': text, 'answer': None}
                else:
                    qa_pairs[qid]['question'] = text
            elif content_match:
                qid = content_match.group(1)
                if qid not in qa_pairs:
                    qa_pairs[qid] = {'question': None, 'answer': text}
                else:
                    qa_pairs[qid]['answer'] = text

    # 过滤掉没有完整问答对的条目（如果需要保留可调整）
    paired = {k: v for k, v in qa_pairs.items() if v['question'] and v['answer']}

    # 根据编号（数字）排序
    sorted_pairs = sorted(paired.items(), key=lambda x: int(x[0]))
    return sorted_pairs

def generate_word_document(sorted_pairs, output_file):
    doc = Document()
    doc.add_heading("FAQ Questions and Answers", 0)

    # 设置文档默认字体为 Times New Roman，字号 12pt
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # 遍历排序后的问答对，依次插入文档（问题一行，答案一行）
    for index, (qid, pair) in enumerate(sorted_pairs, 1):
        question = pair['question']
        answer = pair['answer']
        
        # 去除答案开头的 "Answer:" 或 "Answer：" 或 "Answer " 等
        answer = re.sub(r'^Answer[:：]?\s*', '', answer, flags=re.IGNORECASE)
        
        # 将答案中的数字列表格式从 "1. " 改为 "1) "
        # 匹配行首或换行符后的数字+点+空格的模式
        answer = re.sub(r'(^|\n)(\d+)\.\s+', r'\1\2) ', answer)
        
        doc.add_paragraph(f"{index}. {question}")
        doc.add_paragraph(f"{answer}\n")

    doc.save(output_file)
    print(f"Word document saved as {output_file}")

def main():
    # 请修改此处为你的 Excel 文件路径
    excel_file = r'C:\\Users\\admin\\Desktop\\FAQ\\FAQ_files\\G100.xlsx'
    file_name = os.path.basename(excel_file)
    base_name = os.path.splitext(file_name)[0]
    output_file = f'C:\\Users\\admin\\Desktop\\FAQ\\{base_name}_FAQ.docx'
    
    sorted_pairs = extract_questions_and_answers(excel_file)
    generate_word_document(sorted_pairs, output_file)

if __name__ == '__main__':
    main()