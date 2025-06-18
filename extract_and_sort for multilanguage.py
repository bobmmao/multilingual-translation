import openpyxl
import re
from docx import Document
from docx.shared import Pt
import os

# 待处理的多语种列表
LANGUAGES = [
    "ar-SA", "cs-CZ", "de-DE", "el-GR", "en-US", "es-ES", "fr-FR", "hu-HU",
    "id-ID", "it-IT", "ja-JP", "ko-KR", "nl-NL", "pl-PL", "pt-PT", "ru-RU",
    "sk-SK", "sv-SE", "th-TH", "tr-TR", "vi-VN", "zh-CN", "zh-HK", "zh-TW"
]

def get_language_columns(sheet):
    """
    从第一行中查找所有指定语种的列标题，返回一个字典 { language: col_index }，col_index 为1-based索引。
    """
    lang_cols = {}
    header = list(sheet.iter_rows(min_row=1, max_row=1, values_only=True))[0]
    for idx, cell in enumerate(header, start=1):
        if isinstance(cell, str):
            title = cell.strip()
            if title in LANGUAGES:
                lang_cols[title] = idx
    return lang_cols

def extract_qa_pairs(excel_file, lang_col_index):
    """
    根据匹配规则，从 Excel 文件中提取问答对：
    - 使用第一列（代码）查找包含 "title-数字" 和 "content-数字" 的字段；
    - 对应行中 lang_col_index 列的内容作为文案。
    返回一个排序后的列表，每个元素为 { 'question': ..., 'answer': ... }，排序依据为数字（题号）。
    """
    wb = openpyxl.load_workbook(excel_file)
    sheet = wb.active

    qa_pairs = {}  # 键为题号（字符串），值为 {'question': ..., 'answer': ...}

    # 遍历从第二行开始的所有行
    for row in sheet.iter_rows(min_row=2, values_only=True):
        code_cell = row[0]  # 第一列为代码
        # 因为 values_only 返回的是元组，列索引为 lang_col_index-1
        text_cell = row[lang_col_index - 1]

        if isinstance(code_cell, str) and isinstance(text_cell, str):
            code_str = code_cell.strip()
            text = text_cell.strip()

            # 匹配 title-数字 和 content-数字，注意单元格中只要包含关键词即可
            title_match = re.search(r'title-(\d+)', code_str, re.IGNORECASE)
            content_match = re.search(r'content-(\d+)', code_str, re.IGNORECASE)

            if title_match:
                qid = title_match.group(1).strip()
                if qid not in qa_pairs:
                    qa_pairs[qid] = {'question': text, 'answer': None}
                else:
                    qa_pairs[qid]['question'] = text
            elif content_match:
                qid = content_match.group(1).strip()
                if qid not in qa_pairs:
                    qa_pairs[qid] = {'question': None, 'answer': text}
                else:
                    qa_pairs[qid]['answer'] = text

    # 保留同时有问题和答案的对（如有需要，可保留部分空答案）
    complete_pairs = {k: v for k, v in qa_pairs.items() if v['question'] and v['answer']}
    # 根据题号数字排序
    sorted_pairs = sorted(complete_pairs.items(), key=lambda x: int(x[0]))
    # 返回仅问答列表，不包含题号键
    return [pair for qid, pair in sorted_pairs]

def generate_word_document(qa_list, output_file, language):
    """
    生成 Word 文档，标题为语言标识，每个问答对占两行：一行问题（带序号），一行答案。
    """
    doc = Document()
    doc.add_heading(language, 0)

    # 设置文档默认字体为 Times New Roman，字号12pt
    style = doc.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for idx, pair in enumerate(qa_list, start=1):
        question = pair['question']
        answer = pair['answer']
        doc.add_paragraph(f"{idx}. {question}")
        doc.add_paragraph(f"{answer}\n")

    doc.save(output_file)
    print(f"Word document saved as {output_file}")

def main():
    # Excel 文件路径，请根据实际情况修改
    excel_file = r'C:\\Users\\admin\\Desktop\\FAQ\\FAQ_files\\W100.xlsx'
    
    wb = openpyxl.load_workbook(excel_file)
    sheet = wb.active

    # 获取 Excel 文件中所有语种列（标题在第一行）
    lang_columns = get_language_columns(sheet)
    print("Found language columns:", lang_columns)

    # 从Excel文件名生成基本名称
    base_name = os.path.splitext(os.path.basename(excel_file))[0]
    output_dir = r'C:\Users\admin\Desktop\FAQ'

    # 针对每个指定的语种，只有当该语种存在时进行处理
    for lang in LANGUAGES:
        if lang in lang_columns:
            col_index = lang_columns[lang]
            qa_list = extract_qa_pairs(excel_file, col_index)
            if qa_list:
                output_file = os.path.join(output_dir, f"{base_name}_{lang}.docx")
                generate_word_document(qa_list, output_file, lang)
            else:
                print(f"No complete QA pairs found for {lang}, skipping document generation.")
        else:
            print(f"Language {lang} not found in Excel header, skipping.")

if __name__ == '__main__':
    main()
