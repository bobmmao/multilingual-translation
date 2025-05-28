import os
import docx
import requests
import re
import pandas as pd
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import zipfile
import tempfile
import shutil
from lxml import etree
from tqdm import tqdm
import time

# API Configuration
API_URL = "http://bedroc-proxy-kbhtlwhggrzm-944033383.us-west-2.elb.amazonaws.com/api/v1/chat/completions"
API_KEY = "aws_Bedrock123@pSwD"
MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"

# Target languages
LANGUAGES = {
    "Chinese": "ZH",
}

class TranslationManager:
    def __init__(self):
        self.translation_memory = {}
        self.terminology_db = {}
    
    def load_terminology(self, sheet_url, source_lang_col, target_lang_col, language_code):
        """Load terminology from Google Sheets"""
        try:
            scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
            credentials = ServiceAccountCredentials.from_json_keyfile_name(
                r'C:\\Users\\admin\\Desktop\\多语种说明书翻译\\Extract_glossary.json', scope)
            gc = gspread.authorize(credentials)
            
            sheet_id = sheet_url.split('/d/')[1].split('/')[0]
            worksheet = gc.open_by_key(sheet_id).sheet1
            data = worksheet.get_all_records()
            
            if not data or source_lang_col not in data[0].keys() or target_lang_col not in data[0].keys():
                print(f"Column not found: {source_lang_col} or {target_lang_col}")
                return
            
            term_dict = {str(row[source_lang_col]).strip().lower(): str(row[target_lang_col])
                        for row in data if row[source_lang_col] and row[target_lang_col]}
            
            self.terminology_db[language_code] = term_dict
            print(f"Loaded {len(term_dict)} terms for {language_code}")
            
        except Exception as e:
            print(f"Error loading terminology: {e}")
    
    def apply_terminology(self, text, language_code):
        """Apply terminology to text"""
        if language_code not in self.terminology_db or not self.terminology_db[language_code]:
            return text
        
        terms = sorted(self.terminology_db[language_code].keys(), key=len, reverse=True)
        
        def replace_term(match):
            matched_term = match.group(0)
            replacement = self.terminology_db[language_code][match.group(0).lower()]
            
            if matched_term.islower():
                return replacement.lower()
            elif matched_term.isupper():
                return replacement.upper()
            elif matched_term[0].isupper():
                return replacement[0].upper() + replacement[1:]
            return replacement
        
        for term in terms:
            pattern = r'\b' + re.escape(term) + r'\b'
            text = re.sub(pattern, replace_term, text, flags=re.IGNORECASE)
        
        return text
    
    def collect_context(self, text, language_code, max_examples=3, max_length=300):
        """Collect translation context for similar content"""
        if not self.translation_memory:
            return None
        
        words = re.findall(r'\b\w{4,}\b', text.lower())
        if not words:
            return None
        
        matched_entries = []
        for source, translation in self.translation_memory.items():
            score = sum(1 for word in words if word in source.lower())
            if score > 0:
                matched_entries.append((source, translation, score))
        
        matched_entries.sort(key=lambda x: x[2], reverse=True)
        matched_entries = matched_entries[:max_examples]
        
        if not matched_entries:
            return None
        
        context = []
        for source, translation, _ in matched_entries:
            if len(source) > max_length:
                source = source[:max_length] + "..."
            if len(translation) > max_length:
                translation = translation[:max_length] + "..."
            
            context.append(f"EN: {source}\n{language_code}: {translation}")
        
        return "\n\n".join(context)
    
    def verify_translation(self, translated_text, original_text):
        """Clean and verify translation"""
        english_indicators = [
            "I'm sorry", "I apologize", "Here is the translation",
            "Translated text", "Please note", "I cannot", "I would",
            "Dear Valued Customer", "Best regards", "The Customer Service Team"
        ]
        
        for phrase in english_indicators:
            if phrase.lower() in translated_text.lower():
                clean_lines = [line for line in translated_text.split('\n') 
                              if not any(indicator.lower() in line.lower() 
                                        for indicator in english_indicators)]
                translated_text = '\n'.join(clean_lines)
        
        # Return original if it's just symbols or translation is empty
        if (not re.search(r'[a-zA-Z]', original_text) and len(original_text.strip()) < 5) or \
           (not translated_text.strip() and original_text.strip()):
            return original_text
            
        return translated_text
    
    def translate_text(self, text, target_language, language_code, context=None):
        """Translate text using API with context and terminology support"""
        if not text.strip():
            return ""
        

        # Check translation memory
        memory_key = text.strip().lower()
        if memory_key in self.translation_memory:
            return self.translation_memory[memory_key]
        
        # Extract terminology for context
        potential_terms = []
        if language_code in self.terminology_db:
            for term in self.terminology_db[language_code].keys():
                if re.search(r'\b' + re.escape(term) + r'\b', text, re.IGNORECASE):
                    target_term = self.terminology_db[language_code][term]
                    potential_terms.append(f"{term} -> {target_term}")
        
        try:
            # Skip simple symbols/short text
            if len(text.strip()) < 5 and not re.search(r'[a-zA-Z]', text):
                return text
            
            # Build prompt
            sys_prompt = 'You are a translation engine only. Translate the text to the target language maintaining all formatting. Return ONLY the translated text with no explanations, and no comments. Never apologize or explain your translation.'

            user_prompt = f'Translate the following text to {target_language}. Return ONLY the translated content. Keep all symbols, punctuation, and formatting exactly as they appear. Do not add any explanations, or comments before or after the translation.'
            user_prompt += "\n\nIMPORTANT: If the text contains only symbols, formatting characters, or no text at all (like '----', '***', etc.), do not translate or explain anything - just return those exact symbols."
            user_prompt += "这个是文献翻译，请使中文符合正常的翻译规范，符合文献表达的要求，其中文献的引用要求保留原文不需要翻译，例如 (Li et al., 2022; Shang et al., 2022; Shen et al., 2022)、(Mou, 2020)等。"

            if potential_terms:
                user_prompt += f"\n\nIMPORTANT: Use the following terminology consistently:\n" + "\n".join(
                    potential_terms)

            if context:
                user_prompt += f"\n\nFor consistency, here are some previous translations:\n{context}\n\n"

            user_prompt += f"\nText to translate:\n{text}"

            # Request data
            data = {
                "model": MODEL_ID,
                "messages": [
                    {
                        "role": "user",
                        "content": [{"type": "text", "text": user_prompt}],
                    }
                ],
                "system": sys_prompt
            }
            
            headers = {
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {API_KEY}'
            }
            
            # Send request
            response = requests.post(API_URL, headers=headers, json=data, timeout=60)
            
            if response.status_code == 200:
                result = response.json()
                translated_text = result["choices"][0]["message"]["content"]
                
                # Clean up translation
                translated_text = self.verify_translation(translated_text, text)
                explanation_patterns = [
                    r'^(I\'m sorry|I apologize|Sorry|Note|Please note).*?\n\n',
                    r'\n\n(I\'m sorry|I apologize|Sorry|Note|Please note).*?$',
                    r'^(Here is|Here\'s|The following is|This is) the translation.*?\n\n',
                    r'^Translated text:.*?\n\n'
                ]
                
                for pattern in explanation_patterns:
                    translated_text = re.sub(pattern, '', translated_text, flags=re.IGNORECASE | re.DOTALL)

                # Apply terminology
                translated_text = self.apply_terminology(translated_text, language_code)
                
                # Store in memory
                self.translation_memory[memory_key] = translated_text
                
                return translated_text
            else:
                print(f"Translation error: HTTP {response.status_code} - {response.text}")
                return text
        
        except Exception as e:
            print(f"Translation error: {e}")
            return text

class DocumentProcessor:
    def __init__(self, translator):
        self.translator = translator
    
    def capture_run_properties(self, run):
        """
        Capture all run properties with robust color handling
        """
        properties = {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "strike": run.strike if hasattr(run, "strike") else None,
            "highlight_color": None,  # Initialize as None and only set valid values
            "color": None,  # Initialize as None and only set valid values
            "name": run.font.name if hasattr(run.font, "name") else None,
            "size": run.font.size if hasattr(run.font, "size") else None,
            "subscript": run.font.subscript if hasattr(run.font, "subscript") else None,
            "superscript": run.font.superscript if hasattr(run.font, "superscript") else None,
            "all_caps": run.font.all_caps if hasattr(run.font, "all_caps") else None,
            "double_strike": run.font.double_strike if hasattr(run.font, "double_strike") else None,
            "emboss": run.font.emboss if hasattr(run.font, "emboss") else None,
            "imprint": run.font.imprint if hasattr(run.font, "imprint") else None,
            "outline": run.font.outline if hasattr(run.font, "outline") else None,
            "rtl": run.font.rtl if hasattr(run.font, "rtl") else None,
            "shadow": run.font.shadow if hasattr(run.font, "shadow") else None,
            "small_caps": run.font.small_caps if hasattr(run.font, "small_caps") else None,
            "snap_to_grid": run.font.snap_to_grid if hasattr(run.font, "snap_to_grid") else None,
            "spec_vanish": run.font.spec_vanish if hasattr(run.font, "spec_vanish") else None
        }

        # Handle highlight color with extra caution
        try:
            if hasattr(run.font, "highlight_color"):
                highlight = run.font.highlight_color
                # Only capture if it's a valid value (not None, 'none', or 'auto')
                if highlight is not None and highlight != 'none' and highlight != 'auto':
                    properties["highlight_color"] = highlight
        except Exception:
            pass  # Keep as None if any error

        # Handle font color with extra caution
        try:
            if hasattr(run.font, "color") and run.font.color is not None:
                # Check if a valid RGB value exists
                if hasattr(run.font.color, "rgb") and run.font.color.rgb is not None:
                    # Make sure it's not a problematic value
                    if run.font.color.rgb != 'none' and run.font.color.rgb != 'auto':
                        properties["color"] = run.font.color.rgb
        except Exception:
            pass  # Keep as None if any error
            
        return properties

    def apply_run_properties(self, run, properties):
        """
        Apply run properties with enhanced color handling 
        """
        # Apply basic properties
        run.bold = properties["bold"]
        run.italic = properties["italic"]
        run.underline = properties["underline"]
        
        if properties["strike"] is not None and hasattr(run, "strike"):
            run.strike = properties["strike"]
        
        # Apply font properties
        if properties["name"] is not None:
            run.font.name = properties["name"]
        if properties["size"] is not None:
            run.font.size = properties["size"]
        
        # VERY careful handling of color to avoid 'none' mapping errors
        if properties["color"] is not None and hasattr(run.font, "color"):
            try:
                # Make absolutely sure the color is valid before applying
                color_value = properties["color"]
                if isinstance(color_value, bytes) or (
                    isinstance(color_value, str) and 
                    color_value not in ('none', 'auto', '')
                ):
                    run.font.color.rgb = color_value
            except Exception:
                # If any color error occurs, try to clear the color instead
                try:
                    if hasattr(run.font.color, "_element"):
                        run.font.color._element.clear()
                except:
                    pass
        
        # Equally careful handling of highlight color
        if properties["highlight_color"] is not None and hasattr(run.font, "highlight_color"):
            try:
                # Only apply if it's a valid value
                hl_value = properties["highlight_color"]
                if hl_value not in ('none', 'auto', ''):
                    run.font.highlight_color = hl_value
            except Exception:
                pass
        
        # Apply remaining font properties with error handling
        for attr in ["subscript", "superscript", "all_caps", "double_strike", 
                    "emboss", "imprint", "outline", "rtl", "shadow", 
                    "small_caps", "snap_to_grid", "spec_vanish"]:
            if properties[attr] is not None and hasattr(run.font, attr):
                try:
                    setattr(run.font, attr, properties[attr])
                except Exception:
                    pass
        
        return run

    def process_paragraph(self, paragraph, target_language, language_code):
        """Process and translate paragraph with format preservation"""
        if not paragraph.text.strip():
            return paragraph
        
        # Store original formatting
        runs_formatting = []
        runs_text = []
        
        for run in paragraph.runs:
            runs_text.append(run.text)
            runs_formatting.append(self.capture_run_properties(run))
        
        if not runs_text:
            return paragraph
        
        # Get full paragraph text for translation
        text = paragraph.text
        
        if text.strip():
            # Collect context and translate
            context = self.translator.collect_context(text, language_code)
            translated_text = self.translator.translate_text(text, target_language, language_code, context)
            
            if translated_text == text or not translated_text:
                return paragraph
            
            # Clear paragraph and rebuild with translated content
            paragraph.clear()
            
            # Calculate proportion distributions
            original_total_len = sum(len(run_text) for run_text in runs_text)
            
            if original_total_len == 0:
                run = paragraph.add_run(translated_text)
                if runs_formatting:
                    self.apply_run_properties(run, runs_formatting[0])
            else:
                # Distribute translated text according to original proportions
                run_proportions = [len(run_text) / original_total_len for run_text in runs_text]
                translated_total_len = len(translated_text)
                start_pos = 0
                
                for i, proportion in enumerate(run_proportions):
                    char_count = round(translated_total_len * proportion)
                    end_pos = translated_total_len if i == len(run_proportions) - 1 else min(start_pos + char_count, translated_total_len)
                    
                    run_text = translated_text[start_pos:end_pos]
                    
                    if run_text:
                        new_run = paragraph.add_run(run_text)
                        self.apply_run_properties(new_run, runs_formatting[i])
                    
                    start_pos = end_pos
        
        return paragraph
    
    def process_table(self, table, target_language, language_code):
        """Process and translate table content"""
        try:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        try:
                            self.process_paragraph(para, target_language, language_code)
                        except Exception as e:
                            print(f"Error processing paragraph in table cell: {e}")
        except Exception as e:
            print(f"Error processing table: {e}")
    
    def has_tables(self, doc):
        """Check if document contains tables"""
        return len(doc.tables) > 0
    
    def has_text_boxes(self, doc_path):
        """Check if document contains text boxes"""
        try:
            with zipfile.ZipFile(doc_path, 'r') as zip_ref:
                with tempfile.TemporaryDirectory() as temp_dir:
                    zip_ref.extractall(temp_dir)
                    
                    xml_paths = [
                        os.path.join(temp_dir, 'word/document.xml'),
                        os.path.join(temp_dir, 'word/header*.xml'),
                        os.path.join(temp_dir, 'word/footer*.xml')
                    ]
                    
                    for pattern in xml_paths:
                        for xml_file in glob.glob(pattern) if '*' in pattern else [pattern]:
                            if os.path.exists(xml_file):
                                with open(xml_file, 'r', encoding='utf-8') as f:
                                    if "<w:txbxContent" in f.read():
                                        return True
            return False
        except Exception as e:
            print(f"Error checking for text boxes: {e}")
            return True  # Assume text boxes exist if check fails
    
    def preserve_images(self, input_file, translated_file):
        """Preserve images with exact positioning"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                original_dir = os.path.join(temp_dir, "original")
                translated_dir = os.path.join(temp_dir, "translated")
                
                os.makedirs(original_dir, exist_ok=True)
                os.makedirs(translated_dir, exist_ok=True)
                
                # Extract both documents
                with zipfile.ZipFile(input_file, 'r') as zip_ref:
                    zip_ref.extractall(original_dir)
                
                with zipfile.ZipFile(translated_file, 'r') as zip_ref:
                    zip_ref.extractall(translated_dir)
                
                # Copy media folder
                original_media_dir = os.path.join(original_dir, "word", "media")
                translated_media_dir = os.path.join(translated_dir, "word", "media")
                
                if os.path.exists(original_media_dir):
                    if os.path.exists(translated_media_dir):
                        shutil.rmtree(translated_media_dir)
                    
                    shutil.copytree(original_media_dir, translated_media_dir)
                    print(f"Copied media folder with {len(os.listdir(original_media_dir))} files")
                
                # Process relationships and content types
                self._process_document_relationships(original_dir, translated_dir)
                self._process_document_content_types(original_dir, translated_dir)
                
                # Process document structure to place images
                self._process_document_structure(original_dir, translated_dir)
                
                # Repackage document
                output_temp = translated_file + ".temp.docx"
                
                with zipfile.ZipFile(output_temp, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                    for root, _, files in os.walk(translated_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, translated_dir)
                            zip_out.write(file_path, arcname)
                
                # Replace original output
                if os.path.exists(translated_file):
                    os.remove(translated_file)
                os.rename(output_temp, translated_file)
                
                print("Image processing complete")
                
        except Exception as e:
            print(f"Error processing images: {e}")
            import traceback
            traceback.print_exc()
    
    def _process_document_relationships(self, original_dir, translated_dir):
        """Process document relationships to maintain image references"""
        original_rels_file = os.path.join(original_dir, "word", "_rels", "document.xml.rels")
        translated_rels_file = os.path.join(translated_dir, "word", "_rels", "document.xml.rels")
        
        if not os.path.exists(original_rels_file):
            return
        
        os.makedirs(os.path.dirname(translated_rels_file), exist_ok=True)
        
        try:
            # Extract image relationships
            image_relations = {}
            tree = etree.parse(original_rels_file)
            root = tree.getroot()
            
            for rel in root:
                if all(attr in rel.attrib for attr in ['Id', 'Type', 'Target']):
                    rel_id = rel.attrib['Id']
                    rel_type = rel.attrib['Type']
                    target = rel.attrib['Target']
                    
                    if 'image' in rel_type.lower() or target.startswith('media/'):
                        image_relations[rel_id] = {
                            'type': rel_type,
                            'target': target
                        }
            
            # Update translated document relationships
            if os.path.exists(translated_rels_file):
                tree = etree.parse(translated_rels_file)
                root = tree.getroot()
                
                existing_ids = {rel.attrib['Id'] for rel in root if 'Id' in rel.attrib}
                
                # Add missing image relationships
                for rel_id, rel_info in image_relations.items():
                    if rel_id not in existing_ids:
                        ns = '{http://schemas.openxmlformats.org/package/2006/relationships}'
                        new_rel = etree.SubElement(root, ns + 'Relationship')
                        new_rel.set('Id', rel_id)
                        new_rel.set('Type', rel_info['type'])
                        new_rel.set('Target', rel_info['target'])
                        new_rel.set('TargetMode', 'Internal')
                
                tree.write(translated_rels_file, encoding='UTF-8', xml_declaration=True)
            else:
                # Copy original relationships file if none exists
                shutil.copy2(original_rels_file, translated_rels_file)
                
        except Exception as e:
            print(f"Error processing relationships: {e}")
            if os.path.exists(original_rels_file):
                shutil.copy2(original_rels_file, translated_rels_file)
    
    def _process_document_content_types(self, original_dir, translated_dir):
        """Process content types to include image formats"""
        content_types_file = os.path.join(original_dir, "[Content_Types].xml")
        translated_types_file = os.path.join(translated_dir, "[Content_Types].xml")
        
        if not os.path.exists(content_types_file) or not os.path.exists(translated_types_file):
            return
            
        try:
            # Read original content types
            original_tree = etree.parse(content_types_file)
            original_root = original_tree.getroot()
            
            # Read translated content types
            translated_tree = etree.parse(translated_types_file)
            translated_root = translated_tree.getroot()
            
            # Get existing extensions
            existing_extensions = {default.attrib['Extension'] 
                                for default in translated_root.findall(".//{*}Default") 
                                if 'Extension' in default.attrib}
            
            # Add missing image extensions
            image_extensions = ['png', 'jpeg', 'jpg', 'gif', 'bmp', 'tiff', 'emf', 'wmf']
            modified = False
            
            for default in original_root.findall(".//{*}Default"):
                if 'Extension' in default.attrib:
                    ext = default.attrib['Extension']
                    if ext in image_extensions and ext not in existing_extensions:
                        translated_root.append(default)
                        modified = True
            
            if modified:
                translated_tree.write(translated_types_file, encoding='UTF-8', xml_declaration=True)
                
        except Exception as e:
            print(f"Error processing content types: {e}")
            shutil.copy2(content_types_file, translated_types_file)
    
    def _process_document_structure(self, original_dir, translated_dir):
        """Process document structure to position images correctly"""
        original_doc_xml = os.path.join(original_dir, "word", "document.xml")
        translated_doc_xml = os.path.join(translated_dir, "word", "document.xml")
        
        if not os.path.exists(original_doc_xml) or not os.path.exists(translated_doc_xml):
            return
            
        try:
            with open(original_doc_xml, 'r', encoding='utf-8') as f:
                original_content = f.read()
            
            with open(translated_doc_xml, 'r', encoding='utf-8') as f:
                translated_content = f.read()
            
            # Extract paragraphs from both documents
            original_paragraphs = re.findall(r'<w:p\b[^>]*>.*?</w:p>', original_content, re.DOTALL)
            translated_paragraphs = re.findall(r'<w:p\b[^>]*>.*?</w:p>', translated_content, re.DOTALL)
            
            # Find paragraphs with images
            image_paragraphs = []
            for i, para in enumerate(original_paragraphs):
                if '<w:drawing>' in para:
                    img_refs = re.findall(r'r:embed="(rId\d+)"', para)
                    if img_refs:
                        image_paragraphs.append({
                            'index': i,
                            'paragraph': para,
                            'image_refs': img_refs
                        })
            
            # Find missing images
            missing_images = []
            for img_para in image_paragraphs:
                for ref_id in img_para['image_refs']:
                    if f'r:embed="{ref_id}"' not in translated_content:
                        drawing_match = re.search(
                            f'<w:drawing>.*?r:embed="{ref_id}".*?</w:drawing>', 
                            img_para['paragraph'], 
                            re.DOTALL
                        )
                        if drawing_match:
                            missing_images.append({
                                'ref_id': ref_id,
                                'drawing': drawing_match.group(0),
                                'para_index': img_para['index']
                            })
            
            if missing_images:
                # Position images in translated document
                modified_content = translated_content
                original_para_count = len(original_paragraphs)
                translated_para_count = len(translated_paragraphs)
                
                for img_info in missing_images:
                    # Calculate relative position
                    relative_pos = img_info['para_index'] / original_para_count
                    target_para_index = min(max(0, int(relative_pos * translated_para_count)), 
                                        translated_para_count - 1)
                    
                    # Find target paragraph
                    if target_para_index < len(translated_paragraphs):
                        target_para = translated_paragraphs[target_para_index]
                        para_pos = modified_content.find(target_para)
                        
                        if para_pos != -1:
                            # Insert image paragraph after target
                            para_end = para_pos + len(target_para)
                            img_para = f'<w:p><w:r>{img_info["drawing"]}</w:r></w:p>'
                            modified_content = (
                                modified_content[:para_end] + 
                                img_para + 
                                modified_content[para_end:]
                            )
                        else:
                            # Add to end of document if target not found
                            body_end = modified_content.rfind('</w:body>')
                            if body_end != -1:
                                img_para = f'<w:p><w:r>{img_info["drawing"]}</w:r></w:p>'
                                modified_content = (
                                    modified_content[:body_end] + 
                                    img_para + 
                                    modified_content[body_end:]
                                )
                
                # Save modified document
                with open(translated_doc_xml, 'w', encoding='utf-8') as f:
                    f.write(modified_content)
                    
        except Exception as e:
            print(f"Error processing document structure: {e}")
            import traceback
            traceback.print_exc()
    
    def process_text_boxes(self, doc_path, output_path, target_language, language_code):
        """Process text boxes using direct XML manipulation"""
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Extract document
                with zipfile.ZipFile(doc_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)
                
                # Process XML files that might contain text boxes
                xml_files = ['word/document.xml', 'word/header1.xml', 'word/header2.xml', 'word/header3.xml', 
                            'word/footer1.xml', 'word/footer2.xml', 'word/footer3.xml']
                
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                }
                
                for xml_file in xml_files:
                    xml_path = os.path.join(temp_dir, xml_file)
                    if not os.path.exists(xml_path):
                        continue
                    
                    try:
                        # Parse XML
                        tree = etree.parse(xml_path)
                        root = tree.getroot()
                        
                        # Find text boxes
                        textbox_contents = root.xpath('//w:txbxContent', namespaces=namespaces)
                        
                        if textbox_contents:
                            print(f"Found {len(textbox_contents)} text boxes in {xml_file}")
                            
                            for textbox_content in textbox_contents:
                                # Process paragraphs in text box
                                paragraphs = textbox_content.xpath('.//w:p', namespaces=namespaces)
                                
                                for paragraph in paragraphs:
                                    text_elements = paragraph.xpath('.//w:t', namespaces=namespaces)
                                    paragraph_text = ''.join(text_elem.text or '' for text_elem in text_elements)
                                    
                                    if paragraph_text.strip():
                                        # Get runs with formatting
                                        runs = paragraph.xpath('.//w:r', namespaces=namespaces)
                                        run_texts = []
                                        run_formats = []
                                        
                                        for run in runs:
                                            run_text_elements = run.xpath('.//w:t', namespaces=namespaces)
                                            run_text = ''.join(t.text or '' for t in run_text_elements)
                                            
                                            if run_text:
                                                run_texts.append(run_text)
                                                
                                                # Capture formatting
                                                format_props = {
                                                    'bold': bool(run.xpath('.//w:b', namespaces=namespaces)),
                                                    'italic': bool(run.xpath('.//w:i', namespaces=namespaces)),
                                                    'underline': bool(run.xpath('.//w:u', namespaces=namespaces)),
                                                    'color': run.xpath('.//w:color/@val', namespaces=namespaces)[0] 
                                                            if run.xpath('.//w:color/@val', namespaces=namespaces) else None,
                                                    'font': run.xpath('.//w:rFonts/@ascii', namespaces=namespaces)[0]
                                                            if run.xpath('.//w:rFonts/@ascii', namespaces=namespaces) else None,
                                                    'size': run.xpath('.//w:sz/@val', namespaces=namespaces)[0]
                                                            if run.xpath('.//w:sz/@val', namespaces=namespaces) else None
                                                }
                                                
                                                run_formats.append((run, format_props))
                                        
                                        # Translate paragraph
                                        context = self.translator.collect_context(paragraph_text, language_code)
                                        translated_text = self.translator.translate_text(
                                            paragraph_text, target_language, language_code, context)
                                        
                                        # Apply translation if successful
                                        if translated_text != paragraph_text and translated_text.strip():
                                            original_total_len = sum(len(text) for text in run_texts)
                                            
                                            if original_total_len > 0:
                                                # Distribute translated text based on original proportions
                                                run_proportions = [len(text) / original_total_len for text in run_texts]
                                                translated_total_len = len(translated_text)
                                                start_pos = 0
                                                
                                                for i, (run, _) in enumerate(run_formats):
                                                    proportion = run_proportions[i]
                                                    char_count = round(translated_total_len * proportion)
                                                    
                                                    end_pos = translated_total_len if i == len(run_formats) - 1 \
                                                            else min(start_pos + char_count, translated_total_len)
                                                    
                                                    new_text = translated_text[start_pos:end_pos]
                                                    
                                                    # Update text element
                                                    text_elems = run.xpath('.//w:t', namespaces=namespaces)
                                                    if text_elems:
                                                        text_elems[0].text = new_text
                                                        for text_elem in text_elems[1:]:
                                                            text_elem.text = ""
                                                    
                                                    start_pos = end_pos
                                            elif run_formats:
                                                # Update first run if original is empty
                                                text_elems = run_formats[0][0].xpath('.//w:t', namespaces=namespaces)
                                                if text_elems:
                                                    text_elems[0].text = translated_text
                        
                        # Save updated XML
                        tree.write(xml_path, encoding='UTF-8', xml_declaration=True)
                        
                    except Exception as e:
                        print(f"Error processing {xml_file}: {e}")
                
                # Repackage document
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx_file:
                    for root, _, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, temp_dir)
                            docx_file.write(file_path, arcname)
                
                print(f"Updated document saved to {output_path}")
                
        except Exception as e:
            print(f"Error processing text boxes: {e}")
            # Copy original if failed
            if os.path.exists(doc_path) and output_path != doc_path:
                shutil.copy2(doc_path, output_path)

class DocumentTranslator:
    def __init__(self):
        self.translator = TranslationManager()
        self.processor = DocumentProcessor(self.translator)
    
    def translate_document(self, input_file, output_dir, google_sheet_url=None):
        """Translate document to multiple languages"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Load terminology if provided
        if google_sheet_url:
            for language_name, language_code in LANGUAGES.items():
                self.translator.load_terminology(
                    google_sheet_url, 
                    "English",  # Source language column name
                    language_name,  # Target language column name
                    language_code   # Language code
                )
        
        # Process each target language
        for language_name, language_code in LANGUAGES.items():
            try:
                # Clear translation memory for new language
                self.translator.translation_memory = {}
                
                print(f"\nTranslating to {language_name}...")
                
                # Prepare output file
                base_name = os.path.splitext(os.path.basename(input_file))[0]
                output_file = os.path.join(output_dir, f"{base_name}_{language_code}.docx")
                
                # Part 1: Process standard text
                print("Processing standard text content...")
                
                # Load document
                doc = docx.Document(input_file)
                
                # Translate paragraphs
                print("Processing paragraphs...")
                for para in tqdm(doc.paragraphs, desc="Paragraphs"):
                    try:
                        self.processor.process_paragraph(para, language_name, language_code)
                    except Exception as e:
                        print(f"Error processing paragraph: {e}")
                
                # Translate tables if present
                print("Checking for tables...")
                if self.processor.has_tables(doc):
                    print("Processing tables...")
                    for table in tqdm(doc.tables, desc="Tables"):
                        try:
                            self.processor.process_table(table, language_name, language_code)
                        except Exception as e:
                            print(f"Error processing table: {e}")
                
                # Process headers and footers
                print("Processing headers and footers...")
                try:
                    for section in doc.sections:
                        for para in section.header.paragraphs:
                            self.processor.process_paragraph(para, language_name, language_code)
                        
                        for table in section.header.tables:
                            self.processor.process_table(table, language_name, language_code)
                        
                        for para in section.footer.paragraphs:
                            self.processor.process_paragraph(para, language_name, language_code)
                        
                        for table in section.footer.tables:
                            self.processor.process_table(table, language_name, language_code)
                except Exception as e:
                    print(f"Error processing headers/footers: {e}")
                
                # Save intermediate document
                temp_output_file = output_file + ".temp.docx"
                doc.save(temp_output_file)
                print("Standard text translated and saved to temporary file")
                
                # Part 2: Process text boxes
                print("Checking for text boxes...")
                if self.processor.has_text_boxes(temp_output_file):
                    try:
                        print("Processing text boxes...")
                        self.processor.process_text_boxes(temp_output_file, output_file, language_name, language_code)
                        # Remove temp file
                        if os.path.exists(temp_output_file):
                            os.remove(temp_output_file)
                    except Exception as e:
                        print(f"Error processing text boxes: {e}")
                        # Use temp file as output if failed
                        if os.path.exists(temp_output_file):
                            if os.path.exists(output_file):
                                os.remove(output_file)
                            os.rename(temp_output_file, output_file)
                else:
                    print("No text boxes found, skipping text box processing")
                    # Use temp file as output
                    if os.path.exists(temp_output_file):
                        if os.path.exists(output_file):
                            os.remove(output_file)
                        os.rename(temp_output_file, output_file)
                
                # Preserve images
                try:
                    print("Preserving images...")
                    self.processor.preserve_images(input_file, output_file)
                except Exception as e:
                    print(f"Error preserving images: {e}")
                
                # Print terminology statistics
                if language_code in self.translator.terminology_db:
                    print(f"Terminology statistics ({language_code}):")
                    for term, translation in self.translator.terminology_db[language_code].items():
                        count = sum(1 for source in self.translator.translation_memory.keys() 
                                  if re.search(r'\b' + re.escape(term) + r'\b', source, re.IGNORECASE))
                        if count > 0:
                            print(f"  - '{term}' -> '{translation}': used {count} times")
                
                # Avoid rate limits
                print("Waiting to avoid rate limits...")
                time.sleep(2)
                
            except Exception as e:
                print(f"Error processing language {language_name}: {e}")
                import traceback
                traceback.print_exc()
        
        print("\nAll translations completed!")

def main():
    """Main function"""
    # Direct path parameters
    input_file = r"/Users/mango/Desktop/多语种翻译/10英译中.docx"
    output_dir = r"/Users/mango/Desktop/多语种翻译/test2"
    google_sheet_url = "https://docs.google.com/spreadsheets/d/11B4LNWf27Mt_PvqsyKZYmtxeaLmCBPFSQHiiUyY2IC4/edit?gid=0"
    
    if not os.path.exists(input_file):
        print(f"Error: Input file not found: {input_file}")
        return
    
    # Execute translation
    translator = DocumentTranslator()
    translator.translate_document(input_file, output_dir, google_sheet_url)

if __name__ == "__main__":
    # Import glob for file pattern matching
    import glob
    main()